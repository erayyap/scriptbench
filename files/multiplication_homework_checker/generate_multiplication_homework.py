#!/usr/bin/env python3
"""Generate a playful elementary multiplication homework DOCX."""

from __future__ import annotations

import argparse
import math
import random
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass
class WorkedProblem:
    multiplicand: int
    multiplier: int
    solution_lines: list[str]


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Create a DOCX file filled with column-style multiplication practice."
    )
    parser.add_argument(
        "--count",
        type=int,
        default=10,
        help="Number of standard problems to include (default: 10)",
    )
    parser.add_argument(
        "--seed",
        type=int,
        help="Seed for the random number generator to make output reproducible",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("multiplication_homework.docx"),
        help="Path for the generated DOCX file (default: multiplication_homework.docx)",
    )
    return parser


def add_intro(document: Document) -> None:
    document.add_heading("Magnificent Multiplication Mission!", level=0)
    subtitle = document.add_paragraph(
        "Name: ________________________________      Date: ____________"
    )
    subtitle.style = document.styles["Normal"]
    document.add_paragraph(
        "Today we are stretching our math muscles with some column multiplication. "
        "Show all of your work, circle your answers, and color in a star every time you double-check!"
    )
    bullet_points = [
        "Warm up by skip-counting by 3s aloud before you begin.",
        "Line up the digits carefully so the ones, tens, and hundreds stay best friends.",
        "Use the margin to doodle helpful arrays or quick number lines.",
    ]
    intro_list = document.add_paragraph()
    intro_list.style = "List Bullet"
    intro_list.text = bullet_points[0]
    for point in bullet_points[1:]:
        bullet = document.add_paragraph(style="List Bullet")
        bullet.text = point
    document.add_paragraph()


def _clear_cell(cell) -> None:
    """Remove placeholder paragraphs that python-docx adds to new table cells."""
    for paragraph in list(cell.paragraphs):
        element = paragraph._element  # noqa: SLF001 - direct XML access is required here
        element.getparent().remove(element)


def _set_cell_border(cell, *, top=None, bottom=None, left=None, right=None) -> None:
    """Apply custom borders to a single table cell."""
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.append(tc_pr)

    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge, data in ("top", top), ("bottom", bottom), ("left", left), ("right", right):
        if data is None:
            continue
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        for key, value in data.items():
            border.set(qn(f"w:{key}"), str(value))


def _introduce_student_error(correct_value: int) -> int:
    """Return a nearby incorrect value to mimic a student arithmetic slip."""
    if correct_value == 0:
        return 1

    digits = len(str(abs(correct_value)))
    max_power = max(0, digits - 2)

    for _ in range(10):
        power = random.randint(0, max_power)
        magnitude = random.randint(1, 9) * (10**power)
        if random.random() < 0.5 or correct_value - magnitude <= 0:
            candidate = correct_value + magnitude
        else:
            candidate = correct_value - magnitude
        if candidate != correct_value and candidate > 0:
            return candidate

    # Fallback in case randomness could not find a new value quickly.
    return correct_value + 1


def _build_solution_lines(multiplicand: int, multiplier: int, *, make_wrong: bool) -> list[str]:
    digits = [int(char) for char in reversed(str(multiplier))]
    partials = [multiplicand * digit * (10**index) for index, digit in enumerate(digits)]
    correct_total = multiplicand * multiplier
    displayed_total = (
        _introduce_student_error(correct_total) if make_wrong else correct_total
    )

    if len(partials) == 1:
        width = max(
            len(str(multiplicand)),
            len(str(multiplier)),
            len(str(displayed_total)),
        )
        return [str(displayed_total).rjust(width)]

    partial_lengths = []
    for index, value in enumerate(partials):
        base_length = len(str(value))
        if len(partials) > 1:
            base_length += 1  # reserve space for a leading sign/space
        partial_lengths.append(base_length)

    width = max(
        len(str(multiplicand)),
        len(str(multiplier)),
        len(str(displayed_total)),
        *partial_lengths,
    )

    lines: list[str] = []
    for index, value in enumerate(partials):
        sign = "+" if index == len(partials) - 1 else " "
        text = f"{sign}{value}"
        lines.append(text.rjust(width))

    lines.append("-" * width)
    lines.append(str(displayed_total).rjust(width))
    return lines


def _parse_solution_rows(solution_lines: list[str]) -> list[tuple[str, str, bool]]:
    """Convert formatted solution strings into table-friendly rows.

    Returns a list of tuples containing:
    (operator symbol, digits to display, whether the row needs a top border).
    """

    rows: list[tuple[str, str, bool]] = []
    pending_border = True  # first row sits under the multiplier divider

    for raw in solution_lines:
        stripped = raw.strip()
        if not stripped:
            continue

        if set(stripped) == {"-"}:
            pending_border = True
            continue

        symbol = ""
        digits_text = stripped
        if digits_text and digits_text[0] in "+-":
            symbol = digits_text[0]
            digits_text = digits_text[1:].lstrip()

        rows.append((symbol, digits_text, pending_border))
        pending_border = False

    return rows


def add_problem(
    container,
    multiplicand: int,
    multiplier: int,
    solution_lines: list[str],
    *,
    add_spacing: bool = True,
) -> None:
    parsed_rows = _parse_solution_rows(solution_lines)
    digit_width = max(
        len(str(multiplicand)),
        len(str(multiplier)),
        max((len(row[1]) for row in parsed_rows), default=0),
    )

    total_rows = 2 + len(parsed_rows)
    table = container.add_table(rows=total_rows, cols=digit_width + 1)
    table.style = None
    table.autofit = False

    symbol_column_width = Pt(20)
    digit_column_width = Pt(16)

    for cell in table.columns[0].cells:
        cell.width = symbol_column_width
    for column_index in range(1, digit_width + 1):
        for cell in table.columns[column_index].cells:
            cell.width = digit_column_width

    def populate(cell, text: str, *, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
        _clear_cell(cell)
        paragraph = cell.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(16)

    def fill_row(row_index: int, symbol: str, digits_text: str, *, top_border: bool) -> None:
        populate(table.cell(row_index, 0), symbol)
        characters = list(digits_text)
        padded = [""] * digit_width
        for offset, char in enumerate(reversed(characters)):
            if offset >= digit_width:
                break
            padded[digit_width - 1 - offset] = char

        for column_offset, char in enumerate(padded, start=1):
            populate(table.cell(row_index, column_offset), char)
            if top_border:
                _set_cell_border(
                    table.cell(row_index, column_offset),
                    top={"val": "single", "sz": "12"},
                )

        if top_border:
            _set_cell_border(
                table.cell(row_index, 0),
                top={"val": "single", "sz": "12"},
            )

    fill_row(0, "", str(multiplicand), top_border=False)
    fill_row(1, "x", str(multiplier), top_border=False)

    for index, (symbol, digits_text, needs_border) in enumerate(parsed_rows, start=2):
        fill_row(index, symbol, digits_text, top_border=needs_border)

    if add_spacing:
        container.add_paragraph()


def add_problem_set(
    document: Document,
    problems: list[WorkedProblem],
    *,
    columns: int = 3,
    per_page: int = 9,
) -> None:
    if not problems:
        document.add_heading("Practice Problems", level=1)
        return

    section = document.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    column_width = usable_width // columns

    page_group_size = max(1, per_page)

    for page_index in range(0, len(problems), page_group_size):
        chunk = problems[page_index : page_index + page_group_size]

        if page_index > 0:
            document.add_page_break()
            document.add_heading("Practice Problems (continued)", level=1)
        else:
            document.add_heading("Practice Problems", level=1)

        rows = math.ceil(len(chunk) / columns)
        table = document.add_table(rows=rows, cols=columns)
        table.style = "Table Grid"
        table.autofit = False

        for column in table.columns:
            for cell in column.cells:
                cell.width = column_width

        for index, worked in enumerate(chunk):
            row, column_index = divmod(index, columns)
            cell = table.cell(row, column_index)
            _clear_cell(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            add_problem(
                cell,
                worked.multiplicand,
                worked.multiplier,
                worked.solution_lines,
                add_spacing=False,
            )

        total_cells = rows * columns
        if len(chunk) < total_cells:
            for blank_index in range(len(chunk), total_cells):
                row, column_index = divmod(blank_index, columns)
                _clear_cell(table.cell(row, column_index))


def add_challenge(document: Document, problems: list[WorkedProblem]) -> None:
    document.add_heading("Bonus Challenge", level=1)
    document.add_paragraph(
        "Show what you know! Try these extra brain-boosters and explain how you solved them."
    )

    if not problems:
        document.add_paragraph("(Bonus challenge taking a rest today!)")
    else:
        columns = max(1, len(problems))
        table = document.add_table(rows=1, cols=columns)
        table.style = "Table Grid"
        table.autofit = False
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        column_width = usable_width // columns

        for column in table.columns:
            for cell in column.cells:
                cell.width = column_width

        for index, worked in enumerate(problems):
            cell = table.cell(0, index)
            _clear_cell(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            add_problem(
                cell,
                worked.multiplicand,
                worked.multiplier,
                worked.solution_lines,
                add_spacing=False,
            )

    reflection = document.add_paragraph("Reflection Station: ")
    reflection.add_run(
        "Which problem felt easiest? Which one was the trickiest? Circle the one you want to review in class!"
    ).bold = True
    document.add_paragraph()


def add_closing(document: Document) -> None:
    document.add_heading("Math Motivator", level=1)
    document.add_paragraph(
        "Math Fact of the Day: When you multiply by 10, every digit slides one place to the left and says, \"Zero, tag in!\""
    )
    document.add_paragraph(
        "Finish strong! Star Count: * * * * * (Color a star each time you double-check your work.)"
    )
    document.add_paragraph(
        "Parent Signature: ________________________________  Proud Teacher High-Five: (air high-five!)"
    )


def generate_homework(count: int, output: Path) -> None:
    document = Document()
    add_intro(document)
    practice_pairs = [
        (random.randint(12, 99), random.randint(3, 12))
        for _ in range(count)
    ]
    challenge_pairs = [
        (random.randint(101, 299), random.randint(13, 19)),
        (random.randint(45, 89), random.randint(21, 29)),
    ]

    total_problems = len(practice_pairs) + len(challenge_pairs)
    wrong_target = (
        min(total_problems, max(1, int(round(total_problems * 0.45))))
        if total_problems
        else 0
    )
    wrong_indices = (
        set(random.sample(range(total_problems), wrong_target)) if wrong_target else set()
    )

    worked_practice: list[WorkedProblem] = []
    for index, (multiplicand, multiplier) in enumerate(practice_pairs):
        lines = _build_solution_lines(
            multiplicand,
            multiplier,
            make_wrong=index in wrong_indices,
        )
        worked_practice.append(
            WorkedProblem(multiplicand=multiplicand, multiplier=multiplier, solution_lines=lines)
        )

    worked_challenge: list[WorkedProblem] = []
    for offset, (multiplicand, multiplier) in enumerate(challenge_pairs):
        index = len(practice_pairs) + offset
        lines = _build_solution_lines(
            multiplicand,
            multiplier,
            make_wrong=index in wrong_indices,
        )
        worked_challenge.append(
            WorkedProblem(multiplicand=multiplicand, multiplier=multiplier, solution_lines=lines)
        )

    add_problem_set(document, worked_practice)
    add_challenge(document, worked_challenge)
    add_closing(document)
    footer = document.sections[0].footer.paragraphs[0]
    footer.text = f"Generated on {date.today():%B %d, %Y} - Keep practicing!"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(output)


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    if args.count <= 0:
        raise SystemExit("--count must be a positive integer")

    if args.seed is not None:
        random.seed(args.seed)

    generate_homework(args.count, args.output)
    print(f"Created homework at {args.output}")


if __name__ == "__main__":
    main()
